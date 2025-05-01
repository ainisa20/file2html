from flask import Flask, render_template, request, send_file, url_for
import os
import subprocess
import requests
from werkzeug.utils import secure_filename
import uuid
import base64
import re
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 5000 * 1024 * 1024  # 5000MB

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)


@app.route('/', methods=['GET', 'POST'])
def index():
    download_link = None

    if request.method == 'POST':
        file = request.files.get('file')
        html_url = request.form.get('html_url')

        if file and file.filename != '':
            filename = secure_filename(file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            output_filename = os.path.splitext(filename)[0] + '.md'
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

            file.save(input_path)

            docker_cmd = [
                "docker", "run", "--rm",
                "-v", f"{os.path.abspath(input_path)}:/data/inputfile",
                "-w", "/data",
                "vinsonwang/markitdown:latest",
                "--keep-data-uris",
                "inputfile"
            ]

            try:
                with open(output_path, 'w') as f_out:
                    result = subprocess.run(
                        docker_cmd,
                        stdout=f_out,
                        stderr=subprocess.PIPE,
                        text=True
                    )
                if result.returncode != 0:
                    error_message = result.stderr
                    return render_template('markitdown.html', error=f'处理文件失败：{error_message}')

                download_link = url_for('download_file', filename=output_filename)

            except Exception as e:
                return render_template('markitdown.html', error=f'未知错误: {e}')

        elif html_url and html_url.strip() != '':
            random_id = str(uuid.uuid4())
            input_filename = f"{random_id}.html"
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], input_filename)
            output_filename = f"{random_id}.md"
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

            try:
                response = requests.get(html_url)
                response.raise_for_status()

                with open(input_path, 'w', encoding='utf-8') as f_in:
                    f_in.write(response.text)

                docker_cmd = [
                    "docker", "run", "--rm",
                    "-v", f"{os.path.abspath(input_path)}:/data/inputfile",
                    "-w", "/data",
                    "vinsonwang/markitdown:latest",
                    "--keep-data-uris",
                    "-x", ".html",
                    "inputfile"
                ]

                with open(output_path, 'w') as f_out:
                    result = subprocess.run(
                        docker_cmd,
                        stdout=f_out,
                        stderr=subprocess.PIPE,
                        text=True
                    )

                if result.returncode != 0:
                    error_message = result.stderr
                    return render_template('markitdown.html', error=f'处理HTML URL失败：{error_message}')

                download_link = url_for('download_file', filename=output_filename)

            except requests.RequestException as e:
                return render_template('markitdown.html', error=f'下载HTML失败: {e}')
            except Exception as e:
                return render_template('markitdown.html', error=f'处理HTML失败: {e}')

        else:
            return render_template('markitdown.html', error='请上传文件或填写HTML链接！')

    return render_template('markitdown.html', download_link=download_link)


@app.route('/download/<filename>')
def download_file(filename):
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    return send_file(output_path, as_attachment=True)


import unicodedata

def clean_text(text):
    return ''.join(
        c for c in text
        if c in ('\n', '\r', '\t') or (unicodedata.category(c)[0] != 'C')
    )

@app.route('/convert-docx/<filename>')
def convert_to_docx(filename):
    md_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    docx_filename = os.path.splitext(filename)[0] + '.docx'
    docx_path = os.path.join(app.config['OUTPUT_FOLDER'], docx_filename)

    try:
        with open(md_path, 'r', encoding='utf-8', errors='ignore') as f:
            raw_content = f.read()
            content = clean_text(raw_content)

        doc = Document()

        for line in content.split('\n'):
            img_match = re.search(r'!\[.*?\]\((data:image\/[a-zA-Z]+;base64,[^\)]+)\)', line)
            if img_match:
                base64_data = img_match.group(1)
                header, encoded = base64_data.split(',', 1)
                mime_type = header.split(';')[0].split(':')[1]
                ext = mime_type.split('/')[-1]

                img_data = BytesIO(base64.b64decode(encoded))
                img = Image.open(img_data)
                img_io = BytesIO()
                img.save(img_io, format=ext.upper())
                img_io.seek(0)

                doc.add_picture(img_io, width=Inches(4.5))
            else:
                doc.add_paragraph(line)

        doc.save(docx_path)
        return send_file(docx_path, as_attachment=True)

    except Exception as e:
        return f"转换失败: {e}", 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5009)

